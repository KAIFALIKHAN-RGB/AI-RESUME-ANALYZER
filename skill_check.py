import streamlit as st
import PyPDF2
import docx

st.title("AI Resume Analyzer")
uploaded_file = st.file_uploader("Upload your resume (PDF, DOCX)", type=["pdf", "docx"])
 
if uploaded_file is not None:
    st.success("Resume uploaded successfully!")
    st.write("File Name:", uploaded_file.name)

    file_name = uploaded_file.name.lower()
    text = ""

    if file_name.endswith(".pdf"):
        # Read PDF
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
    elif file_name.endswith(".docx"):
        # Read DOCX
        doc = docx.Document(uploaded_file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        text = "\n".join(full_text)

    text = text.lower()

    skills = ["python", "java", "sql", "html", "css", "machine learning", "data analysis"]
    st.subheader("Skills Found in Resume:")

    found_skills = []
    missing_skills = []
    for skill in skills:
        if skill in text:
            found_skills.append(skill)
        else:
            missing_skills.append(skill)

    if found_skills:
        for skill in found_skills:
            st.success("- " + skill)

    if missing_skills:
        st.subheader("Skills Missing from Resume:")
        for skill in missing_skills:
            st.error("- " + skill)

    else :
        st.warning("No skills found in the resume.")