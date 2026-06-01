import streamlit as st
import PyPDF2
import docx

st.title("AI Resume Analyzer")
uploaded_file = st.file_uploader("Upload your resume (PDF, DOCX)", type=["pdf", "docx"])
 
if uploaded_file is not None:
    st.success("Resume uploaded successfully!")
    st.write("File Name:", uploaded_file.name)

    file_name = uploaded_file.name.lower()
    text = ""

    if file_name.endswith(".pdf"):
        # Read PDF
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
    elif file_name.endswith(".docx"):
        # Read DOCX
        doc = docx.Document(uploaded_file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        text = "\n".join(full_text)

    st.subheader("Extracted Text from Resume:")

    if text:
        st.text_area("Resume Text", text, height=300)
    else:
        st.warning("No text found in the resume.")
    