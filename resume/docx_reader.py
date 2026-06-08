"""
Extracts text content from DOCX resume files.
"""

import docx


def extract_docx_text(uploaded_file):
    """
    Extract all text from a DOCX file, including tables.

    Args:
        uploaded_file: Streamlit UploadedFile object.

    Returns:
        str: Extracted text, or raises ValueError if file is unreadable.
    """
    try:
        doc = docx.Document(uploaded_file)

        full_text = []

        for para in doc.paragraphs:
            full_text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        return "\n".join(full_text)

    except Exception as e:
        raise ValueError(f"Could not read DOCX file: {e}") from e
