import streamlit as st
import PyPDF2
import docx
import json
import google.generativeai as genai

# Set page config for a premium wide-layout experience
st.set_page_config(
    page_title="AI Resume Analyzer & ATS Optimizer",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom premium CSS styling
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;800&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Outfit', sans-serif;
    }
    
    .main-title {
        font-size: 2.8rem;
        font-weight: 800;
        background: linear-gradient(135deg, #6C63FF 0%, #3F3D56 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 0.5rem;
    }
    
    .sub-title {
        font-size: 1.2rem;
        color: #6C757D;
        margin-bottom: 2rem;
    }

    .card {
        background: #FFFFFF;
        padding: 1.5rem;
        border-radius: 16px;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.05);
        border: 1px solid #EAEAEA;
        margin-bottom: 1.5rem;
    }
    
    .score-container {
        display: flex;
        align-items: center;
        justify-content: center;
        background: radial-gradient(circle, #F3F0FF 0%, #E8E3FF 100%);
        border-radius: 50%;
        width: 150px;
        height: 150px;
        margin: auto;
        border: 6px solid #6C63FF;
        box-shadow: 0 8px 24px rgba(108, 99, 255, 0.2);
    }
    
    .score-text {
        font-size: 2.5rem;
        font-weight: 800;
        color: #6C63FF;
    }

    .badge-match {
        display: inline-block;
        background-color: #D4EDDA;
        color: #155724;
        padding: 0.35rem 0.75rem;
        margin: 0.25rem;
        border-radius: 50px;
        font-size: 0.9rem;
        font-weight: 600;
    }

    .badge-missing {
        display: inline-block;
        background-color: #F8D7DA;
        color: #721C24;
        padding: 0.35rem 0.75rem;
        margin: 0.25rem;
        border-radius: 50px;
        font-size: 0.9rem;
        font-weight: 600;
    }
    
    .suggestion-card {
        border-left: 4px solid #FFC107;
        background-color: #FFFDF0;
        padding: 1rem;
        border-radius: 8px;
        margin-bottom: 1rem;
    }
</style>
""", unsafe_allow_html=True)

# Helper function to extract text
def extract_text_from_file(uploaded_file):
    file_name = uploaded_file.name.lower()
    text = ""
    if file_name.endswith('.pdf'):
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
    elif file_name.endswith('.docx'):
        doc = docx.Document(uploaded_file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        text = "\n".join(full_text)
    return text

# Sidebar configuration
st.sidebar.header("⚙️ Configuration")

api_key = st.sidebar.text_input(
    "Gemini API Key",
    value="",
    type="password",
    help="Configure your Google Gemini API Key here."
)

model_choice = st.sidebar.selectbox(
    "Choose Model",
    ["gemini-2.5-flash", "gemini-1.5-flash", "gemini-2.0-flash"],
    index=0
)

st.sidebar.markdown("""
---
### How it works:
1. Upload your resume (PDF or DOCX).
2. Paste the target job description.
3. Click **Analyze Resume**.
4. Gemini AI analyzes the skills gap, resume score, and generates feedback and a tailored cover letter.
""")

# App Header
st.markdown('<div class="main-title">AI Resume Analyzer & ATS Optimizer 🚀</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">Leverage state-of-the-art Gemini AI to audit, score, and optimize your resume for any job description.</div>', unsafe_allow_html=True)

col_input, col_info = st.columns([1, 1])

with col_input:
    st.subheader("📄 Upload Resume")
    uploaded_file = st.file_uploader(
        "Supported formats: PDF, DOCX",
        type=["pdf", "docx"]
    )
    
    st.subheader("💼 Target Job Description")
    job_description = st.text_area(
        "Paste the job description you are targeting...",
        height=250,
        placeholder="We are looking for a Software Engineer with 3+ years of experience in Python, SQL, and AWS..."
    )

with col_info:
    if uploaded_file is None or not job_description:
        st.info("💡 To start, upload your resume and paste a target job description in the left column.")
        # Visual premium card before analysis
        st.markdown("""
        <div class="card">
            <h3>✨ Key Features Available</h3>
            <ul>
                <li><b>ATS Fit Score:</b> Instantly get a match percentage matching standard recruiter systems.</li>
                <li><b>Keyword Gap Analysis:</b> Discover missing technical and soft skills highlighted in the job post.</li>
                <li><b>Bullet Point Tailoring:</b> Actionable suggestions for modifying your descriptions.</li>
                <li><b>Tailored Cover Letter:</b> One-click professional cover letter drafting tailored to the job.</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.success("✅ Inputs ready! Click the button below to start the AI analysis.")

# Analysis trigger button
if uploaded_file is not None and job_description:
    if st.button("🚀 Analyze Resume", use_container_width=True):
        
        if not api_key:
            st.error("Error: Please provide a valid Gemini API Key in the sidebar.")
        else:
            with st.spinner("Analyzing resume against job description using Gemini AI..."):
                try:
                    # Configure API
                    genai.configure(api_key=api_key)
                    
                    # Extract Resume Text
                    resume_text = extract_text_from_file(uploaded_file)
                    
                    if not resume_text.strip():
                        st.error("Could not extract any text from the uploaded file. Please make sure it is not empty or scanned/encrypted.")
                    else:
                        # Construct Prompt
                        prompt = f"""
You are an expert recruiter and applicant tracking system (ATS) optimizer. 
Analyze the provided resume against the target job description.
Please ensure you provide at least 3 detailed tailoring suggestions with specific bullet-point rewrites or section improvement ideas.

Resume Text:
\"\"\"
{resume_text}
\"\"\"

Job Description Text:
\"\"\"
{job_description}
\"\"\"

Provide the analysis strictly in JSON format matching this schema:
{{
  "ats_score": <Integer between 0 and 100>,
  "match_summary": "<A 2-3 sentence overview of the candidate's fit>",
  "matched_skills": [<List of technical or soft skills matched from the job description>],
  "missing_skills": [<List of key skills or keywords present in the job description but missing in the resume>],
  "tailoring_suggestions": [
    {{
      "section": "<e.g. Experience, Projects, Summary>",
      "issue": "<Short description of the mismatch or improvement opportunity>",
      "recommendation": "<Specific rewrite or addition suggestion. Provide at least 3 distinct recommendations.>"
    }}
  ],
  "cover_letter": "<A professionally crafted cover letter draft targeting this role based on the candidate's background>"
}}
"""
                        # Call Gemini Model
                        model = genai.GenerativeModel(
                            model_name=model_choice,
                            generation_config={"response_mime_type": "application/json"}
                        )
                        response = model.generate_content(prompt)
                        
                        # Parse JSON response
                        try:
                            result = json.loads(response.text)
                        except json.JSONDecodeError:
                            # Fallback if JSON mode fails or returns wrapper characters
                            clean_text = response.text.strip()
                            if clean_text.startswith("```json"):
                                clean_text = clean_text[7:]
                            if clean_text.endswith("```"):
                                clean_text = clean_text[:-3]
                            result = json.loads(clean_text)
                            
                        # Show Results
                        st.balloons()
                        st.markdown("---")
                        st.subheader("📊 Analysis Results")
                        
                        # Tabs for cleaner view
                        tab_dashboard, tab_skills, tab_suggestions, tab_cover = st.tabs([
                            "🎯 Dashboard & Score", 
                            "🛠️ Skills Gap Analysis", 
                            "💡 Tailoring Suggestions", 
                            "📝 Tailored Cover Letter"
                        ])
                        
                        with tab_dashboard:
                            col_score, col_summary = st.columns([1, 2])
                            with col_score:
                                st.markdown(f"""
                                <div class="score-container">
                                    <div class="score-text">{result.get('ats_score', 0)}%</div>
                                </div>
                                <h4 style='text-align: center; margin-top: 10px;'>ATS Compatibility Score</h4>
                                """, unsafe_allow_html=True)
                                
                            with col_summary:
                                st.markdown(f"""
                                <div class="card" style='height: 100%;'>
                                    <h3>Recruiter Overview</h3>
                                    <p style='font-size: 1.1rem; line-height: 1.6;'>{result.get('match_summary', '')}</p>
                                </div>
                                """, unsafe_allow_html=True)
                                
                        with tab_skills:
                            col_match, col_miss = st.columns(2)
                            
                            with col_match:
                                st.markdown("### Match Keywords & Skills")
                                matched = result.get('matched_skills', [])
                                if matched:
                                    for ms in matched:
                                        st.markdown(f'<span class="badge-match">{ms}</span>', unsafe_allow_html=True)
                                else:
                                    st.write("No direct key matches found.")
                                    
                            with col_miss:
                                st.markdown("### Missing Keywords & Gaps")
                                missing = result.get('missing_skills', [])
                                if missing:
                                    for mis in missing:
                                        st.markdown(f'<span class="badge-missing">{mis}</span>', unsafe_allow_html=True)
                                else:
                                    st.success("Great! No significant skill gaps detected.")
                                    
                        with tab_suggestions:
                            st.markdown("### Actionable Recommendations")
                            suggestions = result.get('tailoring_suggestions') or result.get('suggestions') or result.get('recommendations') or result.get('tailor_suggestions') or []
                            if suggestions:
                                for sug in suggestions:
                                    if isinstance(sug, dict):
                                        st.markdown(f"""
                                        <div class="suggestion-card">
                                            <strong>Section:</strong> {sug.get('section', 'General')}<br/>
                                            <strong>Issue:</strong> {sug.get('issue', '')}<br/>
                                            <strong>Recommendation:</strong> {sug.get('recommendation', '') or sug.get('suggestion', '')}
                                        </div>
                                        """, unsafe_allow_html=True)
                                    else:
                                        st.markdown(f"""
                                        <div class="suggestion-card">
                                            <strong>Recommendation:</strong> {sug}
                                        </div>
                                        """, unsafe_allow_html=True)
                            else:
                                st.success("Your resume is well aligned. No major tailoring recommendations.")
                                
                        with tab_cover:
                            st.markdown("### Tailored Cover Letter")
                            st.text_area(
                                "Generated Cover Letter (Copy and paste to customize)",
                                value=result.get('cover_letter', ''),
                                height=450
                            )
                            
                        # Debugging section
                        st.markdown("---")
                        with st.expander("🛠️ View Raw JSON Response"):
                            st.json(result)
                            
                except Exception as e:
                    st.error(f"An error occurred during analysis: {str(e)}")
